# resume/editor.py

from docx import Document
import io

def edit_resume_docx(file_stream, missing_keywords):
    """
    Adds missing keywords to the 'Skills' section of a DOCX resume.
    Returns the updated DOCX file as a BytesIO stream.
    """
    try:
        # Reset stream position to the beginning
        file_stream.seek(0)
        doc = Document(file_stream)
        
        skills_paragraph = None
        
        # Find the paragraph containing "Skills" (case-insensitive)
        for para in doc.paragraphs:
            if "skills" in para.text.lower():
                skills_paragraph = para
                break
        
        # Format the new keywords
        keywords_to_add = ", ".join(sorted(list(missing_keywords)))

        if skills_paragraph:
            # Append keywords to the existing skills paragraph
            skills_paragraph.add_run(f", {keywords_to_add}")
        else:
            # If no "Skills" section is found, add a new one
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(keywords_to_add)

        # Save the edited document to a byte stream
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0) # Move cursor to the beginning of the stream
        return doc_io
        
    except Exception as e:
        print(f"Error editing DOCX: {e}")
        return None